"""Post-process REPORT.docx for SmartCura FYP reports (A4).

Fixes applied:
1. Set every table to 100% page width, centered, single-line cell spacing, centered v-align
2. Center every image (figure) and add consistent spacing
3. Style every Figure/Table caption: TNR 11 italic, centered, consistent spacing
4. Make every Heading 1/2/3 "Keep with next" with consistent spacing
5. Apply 0.5" hanging indent to all paragraphs in the References section
6. Force justify on body paragraphs

Run AFTER pandoc has generated REPORT.docx:
    python postprocess.py
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_script_dir = os.path.dirname(os.path.abspath(__file__))
input_path = os.path.join(_script_dir, 'REPORT.docx')

if not os.path.exists(input_path):
    print(f'ERROR: {input_path} not found. Run pandoc first.')
    sys.exit(1)

doc = Document(input_path)


def _set_cell_v_align(cell, val='center'):
    tc_pr = cell._tc.get_or_add_tcPr()
    vAlign = tc_pr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        tc_pr.append(vAlign)
    vAlign.set(qn('w:val'), val)


def _set_cell_no_space(cell):
    """Set paragraph spacing inside a cell to 0pt before/after, single line."""
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _set_table_100pct_centered(table):
    """Set table width to 100% of page width and center it."""
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')
    jc = tblPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        tblPr.append(jc)
    jc.set(qn('w:val'), 'center')


def _set_run_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=(0, 0, 0)):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)


# ── 1. Process every table in the document ────────────────────
table_count = 0
for table in doc.tables:
    table_count += 1
    _set_table_100pct_centered(table)
    for row in table.rows:
        for cell in row.cells:
            _set_cell_no_space(cell)
            _set_cell_v_align(cell, 'center')

print(f'[1/6] Processed {table_count} tables (100% width, centered, single-line, centered v-align)')


# ── 2. Process every image (figure) in the document ──────────
img_count = 0
for shape in doc.inline_shapes:
    p = shape._inline.getparent()
    while p is not None and not p.tag.endswith('}p'):
        p = p.getparent()
    if p is not None:
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            pPr.append(jc)
        jc.set(qn('w:val'), 'center')
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:before'), '120')
        spacing.set(qn('w:after'), '60')
    img_count += 1

print(f'[2/6] Centered {img_count} images')


# ── 3. Style every Figure/Table caption ──────────────────────
# Captions are detected by text content, not style name, because
# pandoc assigns "Body Text" or "Normal" depending on context.
caption_count = 0
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith('Figure ') or text.startswith('Table '):
        # Check if it looks like a caption: "Figure X." or "Table X."
        if (text.startswith('Figure ') and '.' in text[:12]) or \
           (text.startswith('Table ') and '.' in text[:12]):
            caption_count += 1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            # Set font on all runs: TNR 11, italic, grey
            for run in p.runs:
                _set_run_font(run, size=11, italic=True, color=(60, 60, 60))
                # Preserve bold on the "Figure X." / "Table X." part
                if run.text.strip().startswith(('Figure ', 'Table ')):
                    run.font.bold = True
                else:
                    run.font.bold = False

print(f'[3/6] Styled {caption_count} figure/table captions (TNR 11 italic, centered)')


# ── 4. Make every Heading 1/2/3 "Keep with next" ─────────────
h_count = 0
for p in doc.paragraphs:
    style_name = p.style.name
    if style_name in ('Heading 1', 'Heading 2', 'Heading 3'):
        h_count += 1
        p.paragraph_format.keep_with_next = True
        if style_name == 'Heading 1':
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
        elif style_name == 'Heading 2':
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif style_name == 'Heading 3':
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)

print(f'[4/6] Configured {h_count} headings (Keep with next, fixed spacing)')


# ── 5. Apply 0.5" hanging indent to References section ───────
in_references = False
ref_count = 0
for p in doc.paragraphs:
    # Detect the References heading (any heading level)
    if 'References' in p.text and 'Heading' in p.style.name:
        in_references = True
        continue
    # If we hit another heading after References, stop
    if in_references and 'Heading' in p.style.name:
        in_references = False
    if in_references and p.text.strip():
        ref_count += 1
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

print(f'[5/6] Applied 0.5" hanging indent to {ref_count} reference entries')


# ── 6. Force justify on body paragraphs ───────────────────────
just_count = 0
for p in doc.paragraphs:
    if p.style.name == 'Normal' and p.text.strip():
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        just_count += 1

print(f'[6/6] Justified {just_count} body paragraphs')


doc.save(input_path)
print(f'\nSaved to {input_path}')
